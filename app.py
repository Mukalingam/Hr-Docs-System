import streamlit as st
import os
import json
import base64
import pandas as pd
from io import BytesIO
from datetime import datetime
from dotenv import load_dotenv
from db import init_db, authenticate, save_extraction, get_user_extractions, get_extraction_by_id, delete_extraction, get_all_users, add_user, delete_user

load_dotenv()
init_db()

# --- Page Config ---
st.set_page_config(
    page_title="ResumeForge",
    page_icon="🔮",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================
# GLOBAL STYLES
# ============================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500;600&display=swap');

:root {
    --bg-primary: #0a0a0f;
    --bg-card: #12121a;
    --bg-card-hover: #1a1a2e;
    --accent-cyan: #00d4ff;
    --accent-purple: #a855f7;
    --accent-pink: #ec4899;
    --accent-green: #10b981;
    --accent-amber: #f59e0b;
    --accent-red: #ef4444;
    --text-primary: #f0f0f5;
    --text-secondary: #8888aa;
    --border-subtle: #1e1e30;
    --gradient-main: linear-gradient(135deg, #00d4ff 0%, #a855f7 50%, #ec4899 100%);
    --gradient-card: linear-gradient(145deg, #12121a 0%, #1a1a2e 100%);
    --shadow-glow: 0 0 40px rgba(0, 212, 255, 0.08);
}

.stApp { background: var(--bg-primary) !important; font-family: 'Outfit', sans-serif !important; }
.stApp > header { background: transparent !important; }
.block-container { padding-top: 1rem !important; padding-bottom: 2rem !important; max-width: 1200px !important; }
#MainMenu, footer, .stDeployButton { display: none !important; }

/* ===== SIDEBAR ===== */
[data-testid="stSidebar"] {
    background: #0d0d14 !important;
    border-right: 1px solid var(--border-subtle) !important;
}
[data-testid="stSidebar"] * { font-family: 'Outfit', sans-serif !important; }

/* ===== HERO ===== */
.hero-container { text-align: center; padding: 2rem 1rem 1.5rem; position: relative; }
.hero-container::before {
    content: ''; position: absolute; top: -60px; left: 50%; transform: translateX(-50%);
    width: 500px; height: 500px;
    background: radial-gradient(circle, rgba(0,212,255,0.06) 0%, rgba(168,85,247,0.03) 40%, transparent 70%);
    pointer-events: none; z-index: 0;
}
.hero-badge {
    display: inline-flex; align-items: center; gap: 6px;
    background: rgba(0,212,255,0.08); border: 1px solid rgba(0,212,255,0.2);
    border-radius: 50px; padding: 6px 16px; font-size: 0.75rem; font-weight: 500;
    color: var(--accent-cyan); text-transform: uppercase; letter-spacing: 1.5px;
    margin-bottom: 1rem; position: relative; z-index: 1;
}
.hero-title {
    font-family: 'Outfit', sans-serif; font-size: 3rem; font-weight: 800;
    background: var(--gradient-main); -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    background-clip: text; line-height: 1.1; margin: 0.5rem 0; position: relative; z-index: 1;
}
.hero-subtitle {
    font-size: 1rem; color: var(--text-secondary); font-weight: 300;
    max-width: 550px; margin: 0.5rem auto 0; line-height: 1.6; position: relative; z-index: 1;
}

/* ===== CARDS ===== */
.glass-card {
    background: var(--gradient-card); border: 1px solid var(--border-subtle);
    border-radius: 16px; padding: 1.5rem; margin-bottom: 1rem;
    box-shadow: var(--shadow-glow); transition: all 0.3s ease;
}
.glass-card:hover { border-color: rgba(0,212,255,0.15); box-shadow: 0 0 60px rgba(0,212,255,0.06); }
.card-header {
    display: flex; align-items: center; gap: 10px; margin-bottom: 1rem;
    padding-bottom: 0.75rem; border-bottom: 1px solid var(--border-subtle);
}
.card-icon {
    width: 36px; height: 36px; border-radius: 10px;
    display: flex; align-items: center; justify-content: center; font-size: 1.1rem;
}
.icon-cyan { background: rgba(0,212,255,0.1); border: 1px solid rgba(0,212,255,0.2); }
.icon-purple { background: rgba(168,85,247,0.1); border: 1px solid rgba(168,85,247,0.2); }
.icon-pink { background: rgba(236,72,153,0.1); border: 1px solid rgba(236,72,153,0.2); }
.icon-green { background: rgba(16,185,129,0.1); border: 1px solid rgba(16,185,129,0.2); }
.icon-amber { background: rgba(245,158,11,0.1); border: 1px solid rgba(245,158,11,0.2); }
.card-title { font-size: 1rem; font-weight: 600; color: var(--text-primary); margin: 0; }
.card-desc { font-size: 0.78rem; color: var(--text-secondary); margin: 0; }

/* ===== UPLOAD ===== */
[data-testid="stFileUploader"] {
    background: rgba(0,212,255,0.02) !important; border: 2px dashed rgba(0,212,255,0.15) !important;
    border-radius: 12px !important; padding: 0.5rem !important; transition: all 0.3s ease !important;
}
[data-testid="stFileUploader"]:hover { border-color: rgba(0,212,255,0.3) !important; background: rgba(0,212,255,0.04) !important; }

/* ===== BUTTONS ===== */
.stButton > button {
    font-family: 'Outfit', sans-serif !important; font-weight: 600 !important;
    border-radius: 10px !important; padding: 0.6rem 1.5rem !important; transition: all 0.3s ease !important;
}
div[data-testid="stButton"] > button[kind="primary"], .stButton > button[kind="primary"] {
    background: var(--gradient-main) !important; color: white !important;
    border: none !important; box-shadow: 0 4px 20px rgba(0,212,255,0.2) !important;
}
div[data-testid="stButton"] > button[kind="primary"]:hover {
    box-shadow: 0 6px 30px rgba(0,212,255,0.35) !important; transform: translateY(-1px) !important;
}

/* ===== STATUS ===== */
.status-badge {
    display: inline-flex; align-items: center; gap: 6px; padding: 4px 12px;
    border-radius: 50px; font-size: 0.72rem; font-weight: 600; letter-spacing: 0.5px; text-transform: uppercase;
}
.badge-success { background: rgba(16,185,129,0.1); color: #10b981; border: 1px solid rgba(16,185,129,0.2); }
.badge-processing {
    background: rgba(245,158,11,0.1); color: #f59e0b; border: 1px solid rgba(245,158,11,0.2);
    animation: pulse-badge 2s infinite;
}
@keyframes pulse-badge { 0%, 100% { opacity: 1; } 50% { opacity: 0.6; } }

/* ===== DATA ===== */
.data-field { display: flex; justify-content: space-between; align-items: flex-start; padding: 8px 0; border-bottom: 1px solid rgba(30,30,48,0.6); }
.data-field:last-child { border-bottom: none; }
.field-label { font-size: 0.78rem; color: var(--text-secondary); font-weight: 500; min-width: 160px; flex-shrink: 0; }
.field-value { font-size: 0.82rem; color: var(--text-primary); font-weight: 400; text-align: right; word-break: break-word; font-family: 'JetBrains Mono', monospace; }
.section-title { font-family: 'Outfit', sans-serif; font-size: 1.2rem; font-weight: 700; color: var(--text-primary); margin: 1.5rem 0 0.75rem; padding-left: 12px; border-left: 3px solid var(--accent-cyan); }

/* ===== EXPANDER ===== */
[data-testid="stExpander"] { background: var(--bg-card) !important; border: 1px solid var(--border-subtle) !important; border-radius: 12px !important; overflow: hidden; }
[data-testid="stExpander"] summary { color: var(--text-primary) !important; font-family: 'Outfit', sans-serif !important; font-weight: 500 !important; }

/* ===== TABS ===== */
.stTabs [data-baseweb="tab-list"] { gap: 4px; background: var(--bg-card); border-radius: 12px; padding: 4px; border: 1px solid var(--border-subtle); }
.stTabs [data-baseweb="tab"] { border-radius: 8px !important; color: var(--text-secondary) !important; font-family: 'Outfit', sans-serif !important; font-weight: 500 !important; padding: 8px 20px !important; }
.stTabs [aria-selected="true"] { background: rgba(0,212,255,0.1) !important; color: var(--accent-cyan) !important; }

/* ===== INPUTS ===== */
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background: var(--bg-card) !important; border: 1px solid var(--border-subtle) !important;
    border-radius: 10px !important; color: var(--text-primary) !important; font-family: 'Outfit', sans-serif !important;
}
[data-testid="stTextInput"] input:focus, [data-testid="stTextArea"] textarea:focus { border-color: var(--accent-cyan) !important; box-shadow: 0 0 0 1px var(--accent-cyan) !important; }
hr { border: none !important; height: 1px !important; background: var(--border-subtle) !important; margin: 1.5rem 0 !important; }

/* ===== DOWNLOAD ===== */
[data-testid="stDownloadButton"] button { background: rgba(16,185,129,0.1) !important; color: #10b981 !important; border: 1px solid rgba(16,185,129,0.25) !important; border-radius: 10px !important; font-family: 'Outfit', sans-serif !important; font-weight: 600 !important; }
[data-testid="stDownloadButton"] button:hover { background: rgba(16,185,129,0.2) !important; border-color: rgba(16,185,129,0.4) !important; }

.preview-box {
    background: #0d0d14; border: 1px solid var(--border-subtle); border-radius: 12px; padding: 1rem;
    max-height: 400px; overflow-y: auto; font-family: 'JetBrains Mono', monospace;
    font-size: 0.78rem; color: var(--text-secondary); line-height: 1.7; white-space: pre-wrap;
}

.metric-row { display: flex; gap: 12px; margin: 1rem 0; }
.metric-card { flex: 1; background: var(--bg-card); border: 1px solid var(--border-subtle); border-radius: 12px; padding: 12px 16px; text-align: center; }
.metric-value { font-family: 'JetBrains Mono', monospace; font-size: 1.3rem; font-weight: 700; background: var(--gradient-main); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
.metric-label { font-size: 0.7rem; color: var(--text-secondary); text-transform: uppercase; letter-spacing: 1px; margin-top: 2px; }

/* ===== LOGIN ===== */
.login-container {
    max-width: 420px; margin: 6vh auto; padding: 2.5rem;
    background: var(--gradient-card); border: 1px solid var(--border-subtle);
    border-radius: 20px; box-shadow: 0 0 80px rgba(0,212,255,0.06);
}
.login-logo { text-align: center; margin-bottom: 1.5rem; }
.login-logo-text {
    font-size: 2rem; font-weight: 800;
    background: var(--gradient-main); -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.login-subtitle { text-align: center; color: var(--text-secondary); font-size: 0.85rem; margin-bottom: 1.5rem; }

/* ===== HISTORY CARDS ===== */
.history-item {
    background: rgba(18,18,26,0.8); border: 1px solid var(--border-subtle);
    border-radius: 10px; padding: 10px 12px; margin-bottom: 8px;
    cursor: pointer; transition: all 0.2s ease;
}
.history-item:hover { border-color: rgba(0,212,255,0.25); background: rgba(26,26,46,0.8); }
.history-name { font-size: 0.82rem; font-weight: 600; color: var(--text-primary); margin: 0; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
.history-meta { font-size: 0.68rem; color: var(--text-secondary); margin: 2px 0 0; }

/* ===== USER BADGE ===== */
.user-badge {
    display: flex; align-items: center; gap: 10px;
    background: rgba(0,212,255,0.05); border: 1px solid rgba(0,212,255,0.12);
    border-radius: 12px; padding: 10px 14px; margin-bottom: 1rem;
}
.user-avatar {
    width: 34px; height: 34px; border-radius: 8px;
    background: var(--gradient-main); display: flex; align-items: center;
    justify-content: center; font-weight: 700; font-size: 0.85rem; color: white;
}
.user-name { font-size: 0.85rem; font-weight: 600; color: var(--text-primary); margin: 0; }
.user-role { font-size: 0.68rem; color: var(--text-secondary); margin: 0; text-transform: capitalize; }

.app-footer { text-align: center; padding: 2rem 0 1rem; color: var(--text-secondary); font-size: 0.72rem; letter-spacing: 0.5px; }

::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--bg-primary); }
::-webkit-scrollbar-thumb { background: var(--border-subtle); border-radius: 3px; }
</style>
""", unsafe_allow_html=True)


# ============================================================
# UTILITY FUNCTIONS
# ============================================================

def extract_text_from_docx(file_bytes):
    from docx import Document as DocxDocument
    doc = DocxDocument(BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text_from_pdf(file_bytes):
    from PyPDF2 import PdfReader
    reader = PdfReader(BytesIO(file_bytes))
    text = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text.append(t)
    return "\n".join(text)


def extract_text(file_bytes, filename):
    ext = filename.lower().split('.')[-1]
    if ext == 'docx':
        return extract_text_from_docx(file_bytes)
    elif ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    else:
        return "Unsupported file format"


def call_claude_api(api_key, resume_text, format_text):
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)

    system_prompt = """You are an expert HR document processor. Your job is to:
1. Analyze the candidate's resume carefully
2. Understand the required format/template structure
3. Extract ALL relevant information from the resume
4. Map it precisely to the format fields

Return a VALID JSON object (no markdown, no code fences) with these exact keys:

{
    "identification": {
        "name": "Full name from resume",
        "title_position": "Current/target position",
        "pan_number": "PAN number if found, else 'Not Available'",
        "passport_no": "Passport number if found, else 'Not Available'"
    },
    "contact": {
        "mobile": "Phone number",
        "email": "Email address",
        "linkedin": "LinkedIn URL if found, else 'Not Available'"
    },
    "experience_summary": {
        "total_years": "Total years of experience (calculate from work history)",
        "top_skills": [
            {"skill": "Skill name", "years": "Years", "last_used": "Year", "score": "1-4"}
        ]
    },
    "education": {
        "highest_qualification": "Degree name",
        "duration": "Start - End year",
        "college_university": "College/University name",
        "year_of_graduation": "Year",
        "certifications": "Any certifications or 'Not Available'"
    },
    "work_experience": [
        {
            "organization": "Company name",
            "designation": "Current or Previous",
            "duration": "Start - End",
            "projects": [
                {
                    "project_name": "Name",
                    "end_client": "Client name if available",
                    "project_type": "Implementation/Support",
                    "start_date": "Month Year",
                    "end_date": "Month Year or Present",
                    "technology_tools": "Tech stack used",
                    "role": "Role title",
                    "responsibilities": "Key responsibilities as a concise paragraph"
                }
            ]
        }
    ],
    "assessment_summary": {
        "current_location": "City from resume",
        "role_applied_for": "Position from resume",
        "total_it_experience": "Total IT years",
        "relevant_experience": "Relevant years",
        "current_role": "Current role and company",
        "notice_period": "If mentioned, else 'Not Specified'"
    }
}

Be thorough and extract every detail. Calculate experience years from dates. Score skills 1-4 based on depth shown in resume."""

    user_prompt = f"""CANDIDATE RESUME:\n---\n{resume_text}\n---\n\nREQUIRED FORMAT TEMPLATE:\n---\n{format_text}\n---\n\nExtract all information from the resume and map it to the format. Return ONLY valid JSON."""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": user_prompt}],
        system=system_prompt
    )

    response_text = message.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = response_text.split("\n", 1)[1]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
    return json.loads(response_text)


def generate_docx(data):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        return h

    def add_field(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor(0x44, 0x44, 0x66)
        run_value = p.add_run(str(value))
        run_value.font.size = Pt(10)
        return p

    title = doc.add_heading('CV Format', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0F)

    add_heading_styled("1. Identification")
    ident = data.get("identification", {})
    for k in ["name", "title_position", "pan_number", "passport_no"]:
        add_field(k.replace("_", " ").title(), ident.get(k, ""))

    add_heading_styled("2. Contact")
    contact = data.get("contact", {})
    for k in ["mobile", "email", "linkedin"]:
        add_field(k.title(), contact.get(k, ""))

    add_heading_styled("3. Experience and Skill Summary")
    exp_sum = data.get("experience_summary", {})
    add_field("Overall Years of Experience", exp_sum.get("total_years", ""))
    skills = exp_sum.get("top_skills", [])
    if skills:
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Skill Name", "Years of Experience", "Last Used (Year)", "Vendor SME Score (1-4)"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True; run.font.size = Pt(9)
        for skill in skills:
            row = table.add_row()
            row.cells[0].text = str(skill.get("skill", ""))
            row.cells[1].text = str(skill.get("years", ""))
            row.cells[2].text = str(skill.get("last_used", ""))
            row.cells[3].text = str(skill.get("score", ""))

    add_heading_styled("4. Education")
    edu = data.get("education", {})
    for k in ["highest_qualification", "duration", "college_university", "year_of_graduation", "certifications"]:
        add_field(k.replace("_", " ").title(), edu.get(k, ""))

    add_heading_styled("5. Work Experience")
    work = data.get("work_experience", [])
    labels = ["A. Current Organization", "B. Previous Organization", "C. Previous Organization", "D. Previous Organization"]
    for idx, org in enumerate(work):
        lbl = labels[idx] if idx < len(labels) else f"Organization {idx+1}"
        add_heading_styled(f"{lbl}: {org.get('organization', '')}", level=2)
        add_field("Duration", org.get("duration", ""))
        for p_idx, proj in enumerate(org.get("projects", []), 1):
            add_heading_styled(f"Project {p_idx}: {proj.get('project_name', '')}", level=3)
            for pk in ["end_client", "project_type", "start_date", "end_date", "technology_tools", "role"]:
                add_field(pk.replace("_", " ").title(), proj.get(pk, "N/A"))
            p = doc.add_paragraph()
            rl = p.add_run("Responsibilities: "); rl.bold = True; rl.font.size = Pt(10)
            p.add_run(proj.get("responsibilities", "")).font.size = Pt(10)

    add_heading_styled("6. Candidate Assessment Summary")
    assess = data.get("assessment_summary", {})
    for k in ["current_location", "role_applied_for", "total_it_experience", "relevant_experience", "current_role", "notice_period"]:
        add_field(k.replace("_", " ").title(), assess.get(k, ""))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf(data):
    from fpdf import FPDF

    class StyledPDF(FPDF):
        def header(self):
            self.set_font('Helvetica', 'B', 10)
            self.set_text_color(0, 180, 220)
            self.cell(0, 8, 'ResumeForge - Candidate Profile', 0, 1, 'C')
            self.set_draw_color(0, 180, 220)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 7)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f'Generated on {datetime.now().strftime("%d %b %Y")} | Page {self.page_no()}', 0, 0, 'C')

        def section_title(self, title):
            self.set_font('Helvetica', 'B', 12)
            self.set_text_color(26, 26, 46)
            self.set_fill_color(230, 240, 250)
            self.cell(0, 8, title, 0, 1, 'L', True)
            self.ln(2)

        def field(self, label, value):
            self.set_font('Helvetica', 'B', 9)
            self.set_text_color(80, 80, 120)
            self.cell(55, 6, label + ":", 0, 0)
            self.set_font('Helvetica', '', 9)
            self.set_text_color(45, 45, 45)
            self.multi_cell(0, 6, str(value) if value else "N/A")
            self.ln(1)

    pdf = StyledPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 18)
    pdf.set_text_color(10, 10, 15)
    pdf.cell(0, 12, 'CV Format', 0, 1, 'C')
    pdf.ln(4)

    pdf.section_title("1. Identification")
    ident = data.get("identification", {})
    for k in ["name", "title_position", "pan_number", "passport_no"]:
        pdf.field(k.replace("_", " ").title(), ident.get(k, ""))
    pdf.ln(3)

    pdf.section_title("2. Contact")
    contact = data.get("contact", {})
    for k in ["mobile", "email", "linkedin"]:
        pdf.field(k.title(), contact.get(k, ""))
    pdf.ln(3)

    pdf.section_title("3. Experience and Skill Summary")
    exp_sum = data.get("experience_summary", {})
    pdf.field("Overall Experience", exp_sum.get("total_years", ""))
    pdf.ln(2)
    skills = exp_sum.get("top_skills", [])
    if skills:
        pdf.set_font('Helvetica', 'B', 8)
        pdf.set_fill_color(26, 26, 46)
        pdf.set_text_color(255, 255, 255)
        col_w = [55, 40, 40, 45]
        for i, h in enumerate(["Skill", "Years", "Last Used", "Score (1-4)"]):
            pdf.cell(col_w[i], 7, h, 1, 0, 'C', True)
        pdf.ln()
        pdf.set_font('Helvetica', '', 8)
        pdf.set_text_color(45, 45, 45)
        for skill in skills:
            pdf.cell(col_w[0], 6, str(skill.get("skill", ""))[:30], 1, 0)
            pdf.cell(col_w[1], 6, str(skill.get("years", "")), 1, 0, 'C')
            pdf.cell(col_w[2], 6, str(skill.get("last_used", "")), 1, 0, 'C')
            pdf.cell(col_w[3], 6, str(skill.get("score", "")), 1, 0, 'C')
            pdf.ln()
    pdf.ln(3)

    pdf.section_title("4. Education")
    edu = data.get("education", {})
    for k in ["highest_qualification", "duration", "college_university", "year_of_graduation", "certifications"]:
        pdf.field(k.replace("_", " ").title(), edu.get(k, ""))
    pdf.ln(3)

    pdf.section_title("5. Work Experience")
    work = data.get("work_experience", [])
    labels = ["A. Current Organization", "B. Previous Organization", "C. Previous Organization", "D. Previous Organization"]
    for idx, org in enumerate(work):
        lbl = labels[idx] if idx < len(labels) else f"Org {idx+1}"
        pdf.set_font('Helvetica', 'B', 10)
        pdf.set_text_color(0, 150, 200)
        pdf.cell(0, 7, f"{lbl}: {org.get('organization', '')}", 0, 1)
        pdf.field("Duration", org.get("duration", ""))
        for p_idx, proj in enumerate(org.get("projects", []), 1):
            pdf.set_font('Helvetica', 'B', 9)
            pdf.set_text_color(100, 60, 180)
            pdf.cell(0, 6, f"  Project {p_idx}: {proj.get('project_name', '')}", 0, 1)
            for pk in ["end_client", "project_type", "start_date", "end_date", "technology_tools", "role", "responsibilities"]:
                pdf.field(pk.replace("_", " ").title(), proj.get(pk, "N/A"))
            pdf.ln(2)
        pdf.ln(2)

    pdf.section_title("6. Candidate Assessment Summary")
    assess = data.get("assessment_summary", {})
    for k in ["current_location", "role_applied_for", "total_it_experience", "relevant_experience", "current_role", "notice_period"]:
        pdf.field(k.replace("_", " ").title(), assess.get(k, ""))

    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


def render_results(data):
    """Render extraction results — shared between new extraction and history view."""

    ident = data.get("identification", {})
    exp_sum = data.get("experience_summary", {})
    work = data.get("work_experience", [])
    total_projects = sum(len(org.get("projects", [])) for org in work)

    st.markdown(f"""
    <div class="metric-row">
        <div class="metric-card"><div class="metric-value">{exp_sum.get("total_years", "N/A")}</div><div class="metric-label">Years Experience</div></div>
        <div class="metric-card"><div class="metric-value">{len(work)}</div><div class="metric-label">Organizations</div></div>
        <div class="metric-card"><div class="metric-value">{total_projects}</div><div class="metric-label">Projects</div></div>
        <div class="metric-card"><div class="metric-value">{len(exp_sum.get("top_skills", []))}</div><div class="metric-label">Key Skills</div></div>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["📊 Extracted Data", "📥 Download Documents", "🔍 Raw JSON"])

    with tab1:
        # Identification
        st.markdown(f"""<div class="glass-card"><div class="card-header"><div class="card-icon icon-cyan">👤</div><div><p class="card-title">Identification</p></div></div>""", unsafe_allow_html=True)
        for k, v in ident.items():
            st.markdown(f'<div class="data-field"><span class="field-label">{k.replace("_"," ").title()}</span><span class="field-value">{v}</span></div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Contact
        contact = data.get("contact", {})
        st.markdown(f"""<div class="glass-card"><div class="card-header"><div class="card-icon icon-purple">📞</div><div><p class="card-title">Contact Details</p></div></div>""", unsafe_allow_html=True)
        for k, v in contact.items():
            st.markdown(f'<div class="data-field"><span class="field-label">{k.replace("_"," ").title()}</span><span class="field-value">{v}</span></div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Skills
        st.markdown(f"""<div class="glass-card"><div class="card-header"><div class="card-icon icon-green">🎯</div><div><p class="card-title">Skills Assessment</p></div></div>""", unsafe_allow_html=True)
        skills = exp_sum.get("top_skills", [])
        if skills:
            df = pd.DataFrame(skills)
            df.columns = [c.replace("_", " ").title() for c in df.columns]
            st.dataframe(df, use_container_width=True, hide_index=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Education
        edu = data.get("education", {})
        st.markdown(f"""<div class="glass-card"><div class="card-header"><div class="card-icon icon-amber">🎓</div><div><p class="card-title">Education</p></div></div>""", unsafe_allow_html=True)
        for k, v in edu.items():
            st.markdown(f'<div class="data-field"><span class="field-label">{k.replace("_"," ").title()}</span><span class="field-value">{v}</span></div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Work
        for idx, org in enumerate(work):
            org_name = org.get("organization", "Unknown")
            st.markdown(f"""<div class="glass-card"><div class="card-header"><div class="card-icon icon-pink">🏢</div><div><p class="card-title">{org_name}</p><p class="card-desc">{org.get("duration", "")}</p></div></div></div>""", unsafe_allow_html=True)
            for proj in org.get("projects", []):
                with st.expander(f"📁 {proj.get('project_name', 'Project')}"):
                    for pk, pv in proj.items():
                        st.markdown(f"**{pk.replace('_',' ').title()}:** {pv}")

    with tab2:
        st.markdown("""<div class="glass-card"><div class="card-header"><div class="card-icon icon-green">💾</div><div><p class="card-title">Download Generated Documents</p><p class="card-desc">Get the formatted CV in Word or PDF</p></div></div></div>""", unsafe_allow_html=True)
        dl1, dl2 = st.columns(2)
        cname = ident.get("name", "Candidate").replace(" ", "_")
        with dl1:
            st.download_button("📄 Download Word (.docx)", generate_docx(data), f"{cname}_CV_Format.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        with dl2:
            st.download_button("📕 Download PDF (.pdf)", generate_pdf(data), f"{cname}_CV_Format.pdf", "application/pdf", use_container_width=True)

    with tab3:
        st.markdown("""<div class="glass-card"><div class="card-header"><div class="card-icon icon-purple">{ }</div><div><p class="card-title">Raw Extracted JSON</p><p class="card-desc">Complete extracted data in JSON format</p></div></div></div>""", unsafe_allow_html=True)
        st.json(data)


# ============================================================
# LOGIN PAGE
# ============================================================

def render_login():
    st.markdown("""
    <div class="login-container">
        <div class="login-logo"><div class="login-logo-text">🔮 ResumeForge</div></div>
        <div class="login-subtitle">Sign in to continue</div>
    </div>
    """, unsafe_allow_html=True)

    with st.container():
        col_l, col_form, col_r = st.columns([1, 1.2, 1])
        with col_form:
            username = st.text_input("Username", placeholder="Enter username")
            password = st.text_input("Password", type="password", placeholder="Enter password")
            login_btn = st.button("🔐  Sign In", type="primary", use_container_width=True)

            if login_btn:
                if username and password:
                    user = authenticate(username, password)
                    if user:
                        st.session_state['logged_in'] = True
                        st.session_state['user'] = user
                        st.session_state['view_history_id'] = None
                        st.rerun()
                    else:
                        st.error("Invalid credentials. Please try again.")
                else:
                    st.warning("Please enter both username and password.")

            st.markdown("""
            <div style="text-align:center; margin-top:1.5rem; padding:12px; background:rgba(0,212,255,0.04); border:1px solid rgba(0,212,255,0.1); border-radius:10px;">
                <p style="color:var(--text-secondary); font-size:0.75rem; margin:0 0 6px;">Default Credentials</p>
                <p style="color:var(--text-primary); font-size:0.78rem; margin:0; font-family:'JetBrains Mono',monospace;">
                    admin / admin@123<br>
                    recruiter / recruit@123<br>
                    manager / manage@123
                </p>
            </div>
            """, unsafe_allow_html=True)


# ============================================================
# SIDEBAR
# ============================================================

def render_sidebar():
    user = st.session_state['user']
    initials = "".join([w[0] for w in user['full_name'].split()[:2]]).upper()

    with st.sidebar:
        st.markdown(f"""
        <div class="user-badge">
            <div class="user-avatar">{initials}</div>
            <div>
                <p class="user-name">{user['full_name']}</p>
                <p class="user-role">{user['role']}</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # --- Navigation ---
        st.markdown("---")
        nav_col1, nav_col2 = st.columns(2)
        with nav_col1:
            if st.button("🔮 New", use_container_width=True):
                st.session_state['view_history_id'] = None
                st.rerun()
        with nav_col2:
            if st.button("🚪 Logout", use_container_width=True):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()

        st.markdown("---")

        # --- History ---
        st.markdown('<p style="color:var(--text-secondary); font-size:0.75rem; text-transform:uppercase; letter-spacing:1px; margin-bottom:8px;">📂 Recent Extractions</p>', unsafe_allow_html=True)

        extractions = get_user_extractions(user['id'])
        if not extractions:
            st.markdown('<p style="color:var(--text-secondary); font-size:0.8rem; text-align:center; padding:1rem;">No extractions yet</p>', unsafe_allow_html=True)
        else:
            for ext in extractions:
                cname = ext.get('candidate_name', 'Unknown')
                etime = ext.get('created_at', '')
                eid = ext['id']

                col_h, col_d = st.columns([4, 1])
                with col_h:
                    if st.button(f"👤 {cname}", key=f"hist_{eid}", use_container_width=True):
                        st.session_state['view_history_id'] = eid
                        st.rerun()
                with col_d:
                    if st.button("🗑️", key=f"del_{eid}"):
                        delete_extraction(eid)
                        if st.session_state.get('view_history_id') == eid:
                            st.session_state['view_history_id'] = None
                        st.rerun()

                st.markdown(f'<p style="color:var(--text-secondary); font-size:0.68rem; margin:-8px 0 6px 8px;">{etime} · {ext.get("resume_filename","")}</p>', unsafe_allow_html=True)

        # --- Admin Panel ---
        if user['role'] == 'admin':
            st.markdown("---")
            st.markdown('<p style="color:var(--text-secondary); font-size:0.75rem; text-transform:uppercase; letter-spacing:1px; margin-bottom:8px;">⚙️ Admin Panel</p>', unsafe_allow_html=True)

            with st.expander("👥 Manage Users"):
                users = get_all_users()
                for u in users:
                    st.markdown(f"**{u['full_name']}** (`{u['username']}`) — {u['role']}")

                st.markdown("---")
                st.markdown("**Add New User**")
                new_user = st.text_input("Username", key="new_username", placeholder="new_user")
                new_pass = st.text_input("Password", key="new_password", type="password", placeholder="password")
                new_name = st.text_input("Full Name", key="new_fullname", placeholder="John Doe")
                new_role = st.selectbox("Role", ["user", "admin"], key="new_role")
                if st.button("➕ Add User", use_container_width=True):
                    if new_user and new_pass and new_name:
                        if add_user(new_user, new_pass, new_name, new_role):
                            st.success(f"User '{new_user}' created!")
                            st.rerun()
                        else:
                            st.error("Username already exists.")
                    else:
                        st.warning("Fill all fields.")


# ============================================================
# MAIN APP
# ============================================================

def render_main_app():
    api_key = os.getenv("ANTHROPIC_API_KEY", "")

    # Check if viewing history
    history_id = st.session_state.get('view_history_id')
    if history_id:
        record = get_extraction_by_id(history_id)
        if record:
            cname = record.get('candidate_name', 'Unknown')
            st.markdown(f"""
            <div class="hero-container">
                <div class="hero-badge">📂 History</div>
                <div class="hero-title">{cname}</div>
                <div class="hero-subtitle">
                    Extracted on {record.get('created_at', '')} · Resume: {record.get('resume_filename', '')}
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("""<div style="text-align:center; margin-bottom:1rem;"><div class="status-badge badge-success">✓ Saved Extraction</div></div>""", unsafe_allow_html=True)
            data = json.loads(record['extracted_json'])
            render_results(data)
            return
        else:
            st.session_state['view_history_id'] = None

    # --- Normal new extraction flow ---
    st.markdown("""
    <div class="hero-container">
        <div class="hero-badge">⚡ Smart Extraction</div>
        <div class="hero-title">ResumeForge</div>
        <div class="hero-subtitle">
            Upload a candidate resume and format template — automatically extract, map, and generate perfectly formatted documents in seconds.
        </div>
    </div>
    """, unsafe_allow_html=True)

    if not api_key:
        st.markdown("""
        <div class="glass-card" style="border-color: rgba(236,72,153,0.3);">
            <div class="card-header" style="border-bottom: none; margin-bottom: 0;">
                <div class="card-icon icon-pink">⚠️</div>
                <div>
                    <p class="card-title" style="color: #ec4899;">Configuration Missing</p>
                    <p class="card-desc">Add your API key to the <code style="color:#00d4ff; background:rgba(0,212,255,0.08); padding:2px 6px; border-radius:4px;">.env</code> file and restart the app.</p>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""<div class="glass-card"><div class="card-header"><div class="card-icon icon-cyan">📄</div><div><p class="card-title">Candidate Resume</p><p class="card-desc">Upload the candidate's resume (PDF or DOCX)</p></div></div></div>""", unsafe_allow_html=True)
        resume_file = st.file_uploader("Upload Resume", type=["pdf", "docx"], key="resume", label_visibility="collapsed")

    with col2:
        st.markdown("""<div class="glass-card"><div class="card-header"><div class="card-icon icon-pink">📋</div><div><p class="card-title">Format Template</p><p class="card-desc">Upload the required format document (PDF or DOCX)</p></div></div></div>""", unsafe_allow_html=True)
        format_file = st.file_uploader("Upload Format", type=["pdf", "docx"], key="format", label_visibility="collapsed")

    # Previews
    if resume_file or format_file:
        st.markdown('<div class="section-title">📑 Document Preview</div>', unsafe_allow_html=True)
        pc1, pc2 = st.columns(2)
        with pc1:
            if resume_file:
                with st.expander("👤 Resume Preview", expanded=False):
                    rb = resume_file.read(); resume_file.seek(0)
                    rt = extract_text(rb, resume_file.name)
                    st.markdown(f'<div class="preview-box">{rt[:3000]}{"..." if len(rt)>3000 else ""}</div>', unsafe_allow_html=True)
        with pc2:
            if format_file:
                with st.expander("📋 Format Preview", expanded=False):
                    fb = format_file.read(); format_file.seek(0)
                    ft = extract_text(fb, format_file.name)
                    st.markdown(f'<div class="preview-box">{ft[:3000]}{"..." if len(ft)>3000 else ""}</div>', unsafe_allow_html=True)

    # Process button
    st.markdown("")
    _, btn_col, _ = st.columns([1, 2, 1])
    with btn_col:
        can_process = api_key and resume_file and format_file
        process_btn = st.button("🔮  Extract & Generate Documents", type="primary", use_container_width=True, disabled=not can_process)

    if not can_process:
        st.markdown('<div style="text-align:center; padding:1rem;"><p style="color:var(--text-secondary); font-size:0.85rem;">↑ Upload both documents to begin</p></div>', unsafe_allow_html=True)

    if process_btn:
        with st.spinner(""):
            st.markdown("""<div style="text-align:center; padding:1rem;"><div class="status-badge badge-processing">⏳ Processing</div><p style="color:var(--text-secondary); font-size:0.85rem; margin-top:0.5rem;">Reading your documents and extracting data...</p></div>""", unsafe_allow_html=True)

            resume_bytes = resume_file.read(); resume_file.seek(0)
            format_bytes = format_file.read(); format_file.seek(0)
            resume_text = extract_text(resume_bytes, resume_file.name)
            format_text = extract_text(format_bytes, format_file.name)

            try:
                data = call_claude_api(api_key, resume_text, format_text)
                st.session_state['extracted_data'] = data
                st.session_state['processed'] = True

                # Save to history
                candidate_name = data.get("identification", {}).get("name", "Unknown")
                user = st.session_state['user']
                save_extraction(user['id'], candidate_name, resume_file.name, format_file.name, data)

            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
                st.session_state['processed'] = False

    # Show results
    if st.session_state.get('processed') and st.session_state.get('extracted_data'):
        data = st.session_state['extracted_data']
        st.markdown("---")
        st.markdown("""<div style="text-align:center; margin:0.5rem 0 1rem;"><div class="status-badge badge-success">✓ Extraction Complete</div></div>""", unsafe_allow_html=True)
        render_results(data)

    st.markdown('<div class="app-footer">ResumeForge — Built with Streamlit · Extracts. Maps. Generates.</div>', unsafe_allow_html=True)


# ============================================================
# ROUTER
# ============================================================

if st.session_state.get('logged_in'):
    render_sidebar()
    render_main_app()
else:
    render_login()